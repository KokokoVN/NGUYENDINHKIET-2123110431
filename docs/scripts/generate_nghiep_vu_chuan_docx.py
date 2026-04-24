from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "NGHIEP_VU_CHUAN.md"
DOCX_PATH = ROOT / "NGHIEP_VU_CHUAN.docx"


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)


def main() -> None:
    if not MD_PATH.exists():
        raise FileNotFoundError(f"Missing source markdown: {MD_PATH}")

    doc = Document()

    for raw_line in MD_PATH.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("- [ ] "):
            doc.add_paragraph(line[6:].strip(), style="List Bullet")
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue

        add_paragraph(doc, line)

    doc.save(DOCX_PATH)
    print(f"Written: {DOCX_PATH}")


if __name__ == "__main__":
    main()
