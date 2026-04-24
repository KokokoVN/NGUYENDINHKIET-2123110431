from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "NGHIEP_VU_SIEU_CHI_TIET_TIENG_VIET.md"
DOCX_PATH = ROOT / "NGHIEP_VU_SIEU_CHI_TIET_TIENG_VIET.docx"


def set_font(run, size=11, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_toc(paragraph):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "Nhấn chuột phải vào Mục lục -> Update Field -> Update entire table."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    r._r.append(fld_begin)
    r._r.append(instr_text)
    r._r.append(fld_separate)
    r._r.append(txt)
    r._r.append(fld_end)
    set_font(r, 11)


def main():
    if not MD_PATH.exists():
        raise FileNotFoundError(f"Missing file: {MD_PATH}")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = title.add_run("TÀI LIỆU NGHIỆP VỤ SIÊU CHI TIẾT")
    set_font(r1, 22, True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("HỆ THỐNG QUẢN LÝ KHÁCH SẠN")
    set_font(r2, 15, True)

    dt = doc.add_paragraph()
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = dt.add_run(f"Ngày phát hành: {datetime.now().strftime('%d/%m/%Y')}")
    set_font(r3, 11)

    doc.add_page_break()

    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = toc_title.add_run("MỤC LỤC")
    set_font(rt, 16, True)
    toc_p = doc.add_paragraph()
    add_toc(toc_p)

    doc.add_page_break()

    for raw in MD_PATH.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.startswith("# "):
            h = doc.add_heading(level=0)
            rr = h.add_run(line[2:].strip())
            set_font(rr, 17, True)
            continue
        if line.startswith("## "):
            h = doc.add_heading(level=1)
            rr = h.add_run(line[3:].strip())
            set_font(rr, 14, True)
            continue
        if line.startswith("### "):
            h = doc.add_heading(level=2)
            rr = h.add_run(line[4:].strip())
            set_font(rr, 12, True)
            continue
        if line.startswith("- [ ] "):
            p = doc.add_paragraph(line[6:].strip(), style="List Bullet")
            for run in p.runs:
                set_font(run, 11)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            for run in p.runs:
                set_font(run, 11)
            continue

        p = doc.add_paragraph()
        rr = p.add_run(line)
        set_font(rr, 11)

    doc.save(DOCX_PATH)
    print(f"Written: {DOCX_PATH}")


if __name__ == "__main__":
    main()
