from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "NGHIEP_VU_CHUAN.md"
DOCX_PATH = ROOT / "NGHIEP_VU_CHUAN_HOAN_CHINH.docx"


def set_font(run, size=11, bold=False):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_toc_paragraph(paragraph):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_text = OxmlElement("w:t")
    fld_text.text = "Nhan chuot phai vao muc luc -> Update Field -> Update entire table."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r = paragraph.add_run()
    r._r.append(fld_begin)
    r._r.append(instr_text)
    r._r.append(fld_separate)
    r._r.append(fld_text)
    r._r.append(fld_end)
    set_font(r, size=11)


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TAI LIEU NGHIEP VU CHUAN")
    set_font(r, size=24, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("HE THONG QUAN LY KHACH SAN")
    set_font(r2, size=17, bold=True)

    doc.add_paragraph("")
    doc.add_paragraph("")

    info = doc.add_table(rows=6, cols=2)
    info.style = "Table Grid"
    rows = [
        ("Ten tai lieu", "NGHIEP_VU_CHUAN"),
        ("Phien ban", "1.0"),
        ("Ngay cap nhat", datetime.now().strftime("%Y-%m-%d")),
        ("Pham vi", "HotelManagement.Api + hotel-management-web + SQL Server"),
        ("Nguoi tao", "Nhom phat trien"),
        ("Trang thai", "Hoan chinh - San sang ban giao"),
    ]
    for i, (k, v) in enumerate(rows):
        c1 = info.cell(i, 0).paragraphs[0].add_run(k)
        set_font(c1, size=11, bold=True)
        c2 = info.cell(i, 1).paragraphs[0].add_run(v)
        set_font(c2, size=11)

    doc.add_paragraph("")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = note.add_run("Tai lieu nay mo ta nghiep vu chinh thuc dung cho trien khai, kiem thu va ban giao.")
    set_font(rn, size=11)


def add_section_title(doc: Document, text: str, level: int):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    set_font(run, size=16 if level == 1 else 13, bold=True)


def add_normal_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=11)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(6)


def build_body_from_markdown(doc: Document):
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    for raw in lines:
        line = raw.strip()
        if not line:
            doc.add_paragraph("")
            continue

        if line.startswith("# "):
            # Skip top-level title from markdown (already has cover title)
            continue
        if line.startswith("## "):
            add_section_title(doc, line[3:].strip(), 1)
            continue
        if line.startswith("### "):
            add_section_title(doc, line[4:].strip(), 2)
            continue
        if line.startswith("- [ ] "):
            p = doc.add_paragraph(line[6:].strip(), style="List Bullet")
            for run in p.runs:
                set_font(run, size=11)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            for run in p.runs:
                set_font(run, size=11)
            continue

        add_normal_paragraph(doc, line)


def main():
    if not MD_PATH.exists():
        raise FileNotFoundError(f"Missing markdown source: {MD_PATH}")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)

    add_cover(doc)

    doc.add_page_break()

    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = toc_title.add_run("MUC LUC")
    set_font(tr, size=16, bold=True)

    toc_paragraph = doc.add_paragraph()
    add_toc_paragraph(toc_paragraph)

    doc.add_page_break()
    build_body_from_markdown(doc)

    # Add final section for signatures
    doc.add_page_break()
    sign_title = doc.add_paragraph()
    sign_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sign_title.add_run("XAC NHAN BAN GIAO")
    set_font(rs, size=14, bold=True)

    sign_tbl = doc.add_table(rows=2, cols=2)
    sign_tbl.style = "Table Grid"
    left = sign_tbl.cell(0, 0).paragraphs[0].add_run("Ben giao")
    set_font(left, bold=True)
    right = sign_tbl.cell(0, 1).paragraphs[0].add_run("Ben nhan")
    set_font(right, bold=True)
    sign_tbl.cell(1, 0).paragraphs[0].add_run("\n\n(Ky, ghi ro ho ten)")
    sign_tbl.cell(1, 1).paragraphs[0].add_run("\n\n(Ky, ghi ro ho ten)")

    doc.save(DOCX_PATH)
    print(f"Written: {DOCX_PATH}")


if __name__ == "__main__":
    main()
