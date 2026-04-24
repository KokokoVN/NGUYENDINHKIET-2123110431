from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUT = Path(__file__).resolve().parents[1] / "NGHIEP_VU_CHI_TIET_DAY_DU_TIENG_VIET.docx"


def style_run(run, size=11, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading("", level=level)
    r = h.add_run(text)
    style_run(r, size=16 if level == 1 else 13, bold=True)


def add_para(doc: Document, text: str, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, size=11, bold=bold)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(6)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            style_run(run, size=11)


def main():
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("TÀI LIỆU NGHIỆP VỤ CHI TIẾT ĐẦY ĐỦ")
    style_run(r, size=22, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("HỆ THỐNG QUẢN LÝ KHÁCH SẠN")
    style_run(rs, size=15, bold=True)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ri = info.add_run(f"Ngày cập nhật: {datetime.now().strftime('%d/%m/%Y')}")
    style_run(ri, size=11)

    doc.add_page_break()

    add_heading(doc, "1. Mục tiêu và phạm vi", 1)
    add_para(
        doc,
        "Tài liệu mô tả đầy đủ nghiệp vụ vận hành của hệ thống quản lý khách sạn, "
        "bao gồm danh mục, đặt phòng, lưu trú, dịch vụ, thanh toán, hóa đơn, phân quyền và quy tắc ngưng hoạt động theo quan hệ cha/con.",
    )
    add_para(
        doc,
        "Phạm vi triển khai: Backend ASP.NET Core (HotelManagement.Api), Frontend React (hotel-management-web), "
        "Cơ sở dữ liệu SQL Server (HotelManagement).",
    )

    add_heading(doc, "2. Vai trò người dùng", 1)
    add_bullets(
        doc,
        [
            "ADMIN: Toàn quyền quản trị hệ thống và cấu hình dữ liệu.",
            "RECEPTION: Vận hành đặt phòng, check-in/check-out, quản lý danh mục cơ bản.",
            "ACCOUNTANT: Nghiệp vụ thanh toán, hóa đơn, hủy giao dịch (VOID).",
            "MANAGER: Giám sát vận hành và theo dõi báo cáo.",
        ],
    )

    add_heading(doc, "3. Quy trình nghiệp vụ cốt lõi", 1)
    add_bullets(
        doc,
        [
            "Thiết lập danh mục: Khách sạn -> Loại phòng -> Phòng -> Dịch vụ.",
            "Tạo và quản lý khách hàng.",
            "Tạo đặt phòng (Reservation), đảm bảo không trùng lịch.",
            "Check-in tạo Stay; check-out đóng Stay.",
            "Ghi nhận ServiceOrder trong thời gian lưu trú.",
            "Tạo Payment theo Stay (tiền mặt, tự động tính tiền).",
            "Tạo Invoice và xuất PDF hóa đơn.",
        ],
    )

    add_heading(doc, "4. Quy tắc thanh toán mới", 1)
    add_bullets(
        doc,
        [
            "Thanh toán chỉ nhận StayId hợp lệ.",
            "Chỉ cho thanh toán khi Stay đã CHECKED_OUT.",
            "Phương thức thanh toán cố định: CASH (tiền mặt).",
            "Số tiền không nhập tay, hệ thống tự tính = Tiền phòng + Tiền dịch vụ ACTIVE.",
            "Không cho thanh toán trùng khi Stay đã có giao dịch PAID.",
        ],
    )

    add_heading(doc, "5. Quy tắc hóa đơn mới", 1)
    add_bullets(
        doc,
        [
            "Danh sách hóa đơn chỉ hiển thị hóa đơn đã thanh toán thành công (PAID).",
            "Mỗi booking chỉ có một hóa đơn.",
            "Hóa đơn chỉ lập sau khi booking đã CHECKED_OUT.",
            "Hỗ trợ xuất PDF qua endpoint /api/invoices/{id}/pdf.",
        ],
    )

    add_heading(doc, "6. Nghiệp vụ ngưng hoạt động cha/con", 1)
    add_para(
        doc,
        "Khi ngưng đối tượng cha (ví dụ Khách sạn hoặc Loại phòng), hệ thống kiểm tra dữ liệu con đang hoạt động.",
    )
    add_bullets(
        doc,
        [
            "Nếu còn dữ liệu con active: backend trả yêu cầu xác nhận (requiresConfirmation = true).",
            "Frontend hiển thị hộp thoại xác nhận lần 2.",
            "Nếu người dùng đồng ý: gọi lại với cascadeChildren=true để ngưng toàn bộ dữ liệu con.",
            "Ví dụ Hotel có thể ngưng kèm RoomType, Room, HotelService.",
        ],
    )

    add_heading(doc, "7. Danh sách API chính", 1)
    add_bullets(
        doc,
        [
            "Auth: /api/auth/login, /api/auth/me",
            "Hotels: /api/hotels, /api/hotels/{id}, /api/hotels/{id}/restore",
            "RoomTypes: /api/roomtypes, /api/roomtypes/{id}, /api/roomtypes/{id}/restore",
            "Rooms: /api/rooms, /api/rooms/available, /api/rooms/{id}/restore",
            "Customers: /api/customers",
            "Bookings/Stays: /api/bookings, /api/stays",
            "Services: /api/hotelservices, /api/service-orders",
            "Payments/Invoices: /api/payments, /api/invoices, /api/invoices/{id}/pdf",
        ],
    )

    add_heading(doc, "8. Checklist nghiệm thu", 1)
    add_bullets(
        doc,
        [
            "Tạo dữ liệu danh mục và đặt phòng thành công.",
            "Check-in/check-out đúng luồng và đúng trạng thái.",
            "Payment chỉ thanh toán được sau check-out.",
            "Invoice list chỉ trả về hóa đơn đã PAID.",
            "Xuất PDF hóa đơn tải file thành công.",
            "Ngưng cha có cảnh báo xác nhận và cascade con hoạt động đúng.",
        ],
    )

    add_heading(doc, "9. Kết luận", 1)
    add_para(
        doc,
        "Tài liệu này là phiên bản tiếng Việt đầy đủ dấu, có thể dùng trực tiếp cho bàn giao, "
        "nghiệm thu hoặc nộp báo cáo đồ án.",
    )

    doc.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    main()
