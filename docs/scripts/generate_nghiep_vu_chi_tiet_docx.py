from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "NGHIEP_VU_CHI_TIET_DAY_DU.md"
DOCX_PATH = ROOT / "NGHIEP_VU_CHI_TIET_DAY_DU.docx"


def set_font(run, size=11, bold=False):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_toc(paragraph):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Right click Muc luc -> Update Field -> Update entire table."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    r._r.append(fld_begin)
    r._r.append(instr_text)
    r._r.append(fld_separate)
    r._r.append(text)
    r._r.append(fld_end)
    set_font(r, 11)


def main():
    if not MD_PATH.exists():
        raise FileNotFoundError(MD_PATH)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TAI LIEU NGHIEP VU CHI TIET DAY DU")
    set_font(r, 22, True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("HE THONG QUAN LY KHACH SAN")
    set_font(r2, 15, True)
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = info.add_run(f"Ngay tao: {datetime.now().strftime('%Y-%m-%d')}")
    set_font(r3, 11)

    doc.add_page_break()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("MUC LUC")
    set_font(tr, 16, True)
    toc_p = doc.add_paragraph()
    add_toc(toc_p)

    doc.add_page_break()
    for raw in MD_PATH.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            h = doc.add_heading(level=0)
            rr = h.add_run(line[2:].strip())
            set_font(rr, 17, True)
            continue
        if line.startswith("## "):
            h = doc.add_heading(level=1)
            rr = h.add_run(line[3:].strip())
            set_font(rr, 14, True)
            continue
        if line.startswith("### "):
            h = doc.add_heading(level=2)
            rr = h.add_run(line[4:].strip())
            set_font(rr, 12, True)
            continue
        if line.startswith("- [ ] "):
            p = doc.add_paragraph(line[6:].strip(), style="List Bullet")
            for run in p.runs:
                set_font(run, 11)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            for run in p.runs:
                set_font(run, 11)
            continue

        p = doc.add_paragraph()
        rr = p.add_run(line)
        set_font(rr, 11)

    doc.save(DOCX_PATH)
    print(f"Written: {DOCX_PATH}")


if __name__ == "__main__":
    main()
