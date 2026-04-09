# -*- coding: utf-8 -*-
"""Sinh file Word mô tả nghiệp vụ hệ thống Quản lý khách sạn."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parents[1] / "TaiLieuNghiepVu_QuanLyKhachSan.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for x in items:
        doc.add_paragraph(x, style="List Bullet")


def main():
    doc = Document()
    title = doc.add_heading("TÀI LIỆU NGHIỆP VỤ", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("Hệ thống Quản lý khách sạn (Hotel Management)\n")
    r.bold = True
    r.font.size = Pt(14)
    st.add_run("Phiên bản tài liệu: 1.0 — Bám sát mã nguồn API ASP.NET Core + Web React\n")
    st.add_run("Ngày lập: tháng 4/2026")

    add_heading(doc, "1. Mục đích và phạm vi", 1)
    add_para(
        doc,
        "Tài liệu mô tả đầy đủ nghiệp vụ của hệ thống quản lý khách sạn: quản lý danh mục, "
        "khách hàng, đặt phòng, lưu trú, dịch vụ phòng, hóa đơn, thanh toán, vận hành phụ trợ "
        "(buồng phòng, bảo trì), báo cáo, phân quyền người dùng và nhật ký kiểm toán.",
    )
    add_para(
        doc,
        "Phạm vi gồm: backend REST API (HotelManagement.Api), ứng dụng web quản trị (hotel-management-web), "
        "cơ sở dữ liệu SQL Server theo mô hình domain đã triển khai.",
    )

    add_heading(doc, "2. Đối tượng sử dụng và vai trò", 1)
    add_para(doc, "Hệ thống phân quyền theo mã vai trò (Role) gắn với tài khoản nội bộ:", bold=True)
    add_bullets(
        doc,
        [
            "ADMIN: Toàn quyền cấu hình, người dùng, vai trò, điều chỉnh tích điểm, hủy thanh toán, v.v.",
            "RECEPTION (Lễ tân): Danh mục khách sạn/loại phòng/phòng, khách, đặt phòng, nhận/trả phòng, một số dịch vụ và hóa đơn.",
            "ACCOUNTANT (Kế toán): Hóa đơn, thanh toán, hủy giao dịch thanh toán (theo quyền API).",
            "MANAGER: Báo cáo, xem nhật ký audit (theo API).",
            "HOUSEKEEPING / MAINTENANCE: Công việc buồng phòng, phiếu bảo trì (theo API).",
        ],
    )
    add_para(
        doc,
        "Đăng nhập: JWT Bearer. Người dùng không có token hợp lệ không gọi được các API được bảo vệ.",
    )

    add_heading(doc, "3. Tổng quan chức năng theo module", 1)
    modules = [
        (
            "Xác thực (Auth)",
            "Đăng nhập, xem thông tin tài khoản hiện tại (/me), đăng xuất phía client (JWT stateless).",
        ),
        (
            "Danh mục vận hành",
            "Khách sạn (Hotel), Loại phòng (RoomType), Phòng (Room): CRUD, xóa mềm theo quy tắc nghiệp vụ.",
        ),
        (
            "Khách hàng (Customer)",
            "Quản lý hồ sơ khách (cá nhân / công ty / đại lý), tra cứu, xóa mềm; tích điểm và hạng thành viên.",
        ),
        (
            "Khách lưu trú (Guest)",
            "Ghi nhận khách đi kèm / thông tin khách trên đơn (theo API).",
        ),
        (
            "Đặt phòng (Booking / Reservation)",
            "Tạo đơn, lọc theo khách sạn/phòng/khách/trạng thái/ngày; nhận phòng, trả phòng, hủy đơn.",
        ),
        (
            "Lưu trú (Stay)",
            "Theo dõi kỳ ở thực tế sau check-in; tra cứu theo đơn/trạng thái.",
        ),
        (
            "Dịch vụ danh mục & đặt dịch vụ",
            "HotelService (danh mục dịch vụ); ServiceOrder (đặt dịch vụ theo Stay hoặc Reservation, hủy dòng dịch vụ).",
        ),
        (
            "Hóa đơn & thanh toán",
            "Xuất hóa đơn sau trả phòng; ghi nhận thanh toán theo Stay hoặc Reservation; void giao dịch PAID.",
        ),
        (
            "Buồng phòng & bảo trì",
            "HousekeepingTask: công việc dọn/phòng; MaintenanceTicket: sự cố / bảo trì phòng.",
        ),
        (
            "Báo cáo",
            "Dashboard tổng hợp phòng, đặt phòng, doanh thu hóa đơn, top dịch vụ.",
        ),
        (
            "Nhật ký & quản trị",
            "AuditLog: ghi nhận thao tác quan trọng và có thể ghi log gọi API; Users/Roles: quản lý tài khoản và vai trò (ADMIN).",
        ),
    ]
    for name, desc in modules:
        p = doc.add_paragraph(style="List Number")
        p.add_run(name + ": ").bold = True
        p.add_run(desc)

    add_heading(doc, "4. Quy trình nghiệp vụ cốt lõi", 1)

    add_heading(doc, "4.1. Thiết lập danh mục", 2)
    add_bullets(
        doc,
        [
            "Tạo Khách sạn đang hoạt động (IsActive).",
            "Tạo Loại phòng thuộc đúng khách sạn: sức chứa, giá cơ bản (base rate).",
            "Tạo Phòng: số phòng duy nhất trong khách sạn (trong các phòng active); trạng thái phòng (VACANT, RESERVED, OCCUPIED, DIRTY, OUT_OF_SERVICE, MAINTENANCE, …).",
            "Xóa mềm phòng: đưa phòng về OUT_OF_SERVICE; xóa mềm loại phòng chỉ khi không còn phòng active dùng loại đó; xóa mềm khách sạn khi không còn phòng/loại phòng active phụ thuộc.",
        ],
    )

    add_heading(doc, "4.2. Quản lý khách hàng", 2)
    add_bullets(
        doc,
        [
            "Khách INDIVIDUAL: bắt buộc họ tên.",
            "Khách COMPANY / AGENCY: bắt buộc tên công ty.",
            "Tra cứu theo từ khóa (họ tên, công ty, SĐT, email).",
            "Xóa mềm: đánh dấu DeletedAt, không hiển thị trong danh sách chuẩn.",
            "Tích điểm: khi xuất hóa đơn cho đặt phòng có CustomerId, hệ thống cộng điểm (quy tắc: 1 điểm / 100.000đ, làm tròn floor); hạng BRONZE / SILVER / GOLD / PLATINUM.",
            "ADMIN có thể điều chỉnh điểm thủ công kèm lý do (ghi audit).",
        ],
    )

    add_heading(doc, "4.3. Đặt phòng (Reservation)", 2)
    add_bullets(
        doc,
        [
            "Tạo đơn: chọn phòng, khách (customerId) hoặc tạo khách mới trong cùng request; ngày nhận/trả; số người; có thể ghi đơn giá đêm hoặc để null lấy theo base rate loại phòng.",
            "Ràng buộc: không trùng lịch với đơn CONFIRMED hoặc CHECKED_IN; phòng OUT_OF_SERVICE / MAINTENANCE không cho đặt; tổng người không vượt capacity loại phòng.",
            "Trạng thái đơn tiêu biểu: CONFIRMED → (check-in) CHECKED_IN → (check-out) CHECKED_OUT; có thể CANCELLED khi còn CONFIRMED.",
        ],
    )

    add_heading(doc, "4.4. Nhận phòng (Check-in)", 2)
    add_bullets(
        doc,
        [
            "Chỉ check-in khi đơn đang CONFIRMED.",
            "Không check-in lặp nếu đã có Stay IN_HOUSE cho đơn đó.",
            "Kết quả: tạo bản ghi Stay (IN_HOUSE), cập nhật trạng thái đơn và phòng (ví dụ OCCUPIED).",
        ],
    )

    add_heading(doc, "4.5. Dịch vụ trong lưu trú", 2)
    add_bullets(
        doc,
        [
            "Đặt dịch vụ khi đã check-in và còn Stay IN_HOUSE (hoặc theo reservationId tùy luồng API).",
            "Mã dịch vụ ví dụ: MINIBAR, LAUNDRY, BREAKFAST, EXTRA_BED, OTHER.",
            "Có thể hủy dòng dịch vụ (cancel) kèm lý do.",
        ],
    )

    add_heading(doc, "4.6. Trả phòng (Check-out)", 2)
    add_bullets(
        doc,
        [
            "Đóng Stay, cập nhật đơn sang CHECKED_OUT, phòng thường chuyển DIRTY để buồng phòng xử lý.",
            "Cần có Stay hợp lệ từ luồng check-in API; dữ liệu cũ thiếu Stay có thể cần xử lý dữ liệu thủ công.",
        ],
    )

    add_heading(doc, "4.7. Hóa đơn", 2)
    add_bullets(
        doc,
        [
            "Chỉ xuất hóa đơn khi đơn đã CHECKED_OUT.",
            "Mỗi đặt phòng chỉ một hóa đơn (không trùng).",
            "Tiền phòng = số đêm × RatePerNight; tiền dịch vụ = tổng (quantity × unitPrice) các ServiceOrder ACTIVE của mọi Stay thuộc đơn.",
            "Ghi nhận phương thức thanh toán, ghi chú, thời điểm thanh toán trên hóa đơn.",
        ],
    )

    add_heading(doc, "4.8. Thanh toán (Payment)", 2)
    add_bullets(
        doc,
        [
            "Ghi nhận thanh toán gắn với đúng một trong hai: stayId hoặc reservationId.",
            "Trạng thái giao dịch: PAID; chỉ cho phép void (chuyển VOID) khi đang PAID, do ADMIN/ACCOUNTANT.",
        ],
    )

    add_heading(doc, "4.9. Buồng phòng & bảo trì", 2)
    add_bullets(
        doc,
        [
            "Housekeeping: tạo/cập nhật/hủy công việc theo phòng và trạng thái công việc.",
            "Maintenance: phiếu sự cố, phân công, cập nhật trạng thái, đóng phiếu.",
        ],
    )

    add_heading(doc, "4.10. Báo cáo", 2)
    add_para(
        doc,
        "Dashboard: thống kê phòng (tổng active, trống, đang có khách, theo status), "
        "đặt phòng theo trạng thái, doanh thu hóa đơn, top dịch vụ trong kỳ.",
    )

    add_heading(doc, "5. Giao diện Web (React)", 1)
    add_bullets(
        doc,
        [
            "Đăng nhập JWT; lưu token và meta người dùng.",
            "Trang: Tổng quan (dashboard), Khách sạn, Loại phòng, Phòng, Khách hàng, Đặt phòng, Hóa đơn, Người dùng (ADMIN, chỉ xem danh sách nếu chưa mở rộng UI).",
            "Các module còn lại (thanh toán, guest, stays, dịch vụ, buồng phòng, bảo trì, audit) chủ yếu qua API/Postman hoặc có thể mở rộng web sau.",
        ],
    )

    add_heading(doc, "6. API, bảo mật và chuẩn phản hồi", 1)
    add_bullets(
        doc,
        [
            "Base path: /api/...",
            "Đăng nhập: POST /api/auth/login; các API khác: header Authorization: Bearer <accessToken>.",
            "CORS cho origin dev frontend (localhost:5173).",
            "Phản hồi thành công có thể được bọc dạng { message, data } theo filter toàn cục; lỗi có message/traceId hoặc chi tiết validate.",
            "Ghi log gọi API và/hoặc ghi AuditLog (API_REQUEST, API_EXCEPTION, và các action nghiệp vụ HOTEL_*, BOOKING_*, …).",
        ],
    )

    add_heading(doc, "7. Ràng buộc và quy tắc quan trọng (tóm tắt)", 1)
    add_bullets(
        doc,
        [
            "Phòng phải khớp khách sạn và loại phòng.",
            "Đặt phòng tuân thủ lịch và trạng thái phòng.",
            "Hóa đơn sau trả phòng; tích điểm khi có khách trên đơn.",
            "Thanh toán void chỉ khi PAID.",
            "Phân quyền theo role trên từng endpoint.",
        ],
    )

    add_heading(doc, "8. Phụ lục — Danh sách nhóm API chính", 1)
    add_para(doc, "Tham chiếu nhanh (không thay thế Swagger):", bold=True)
    apis = [
        "POST /api/auth/login, GET /api/auth/me, POST /api/auth/logout",
        "GET|POST|PUT|DELETE /api/hotels, /api/roomtypes, /api/rooms",
        "GET|POST|PUT|DELETE /api/customers (+ loyalty, adjust)",
        "GET|POST|PUT|DELETE /api/guests",
        "GET|POST /api/bookings; PUT .../check-in, /check-out, /cancel",
        "GET /api/stays",
        "GET|POST|PUT /api/hotelservices",
        "GET|POST /api/serviceorders; PUT .../cancel",
        "GET|POST /api/invoices",
        "GET|POST /api/payments; PUT .../void",
        "GET|POST|PUT|DELETE /api/housekeepingtasks",
        "GET|POST|PUT|DELETE /api/maintenancetickets",
        "GET /api/reports/dashboard",
        "GET /api/auditlogs",
        "GET|POST|PUT /api/users (+ /status, /roles)",
        "GET|POST|PUT /api/roles",
    ]
    for a in apis:
        doc.add_paragraph(a, style="List Bullet")

    add_heading(doc, "9. Kết luận", 1)
    add_para(
        doc,
        "Tài liệu tổng hợp nghiệp vụ đầy đủ theo hệ thống hiện tại, phục vụ giảng dạy, bàn giao, "
        "kiểm thử và mở rộng tính năng. Chi tiết tham số request/response nên đối chiếu Swagger và mã nguồn DTO.",
    )

    doc.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    main()
